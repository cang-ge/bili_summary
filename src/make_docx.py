"""
Redesigned Word (.docx) version with the same editorial design.
Uses Microsoft YaHei for Chinese, Calibri for English.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\Users\Administrator\Desktop\雅思自学流程-Kasa_ZYY-总结.docx'

# Design tokens
INK = RGBColor(0x0F, 0x17, 0x2A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
SLATE = RGBColor(0x64, 0x74, 0x8B)
BORDER = RGBColor(0xE8, 0xE2, 0xD5)
PAPER = RGBColor(0xFA, 0xF8, 0xF5)
EMERALD = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)
AMBER = RGBColor(0xD9, 0x77, 0x06)
BG_DARK = '0F172A'
BG_LIGHT = 'FCFBF8'
BG_BAND = 'F4F0E8'

doc = Document()

# Margins
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_h1(num, title, eyebrow=None):
    if eyebrow:
        p_ = doc.add_paragraph()
        run = p_.add_run(eyebrow.upper())
        run.font.size = Pt(9)
        run.font.color.rgb = GOLD
        run.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p_.paragraph_format.space_after = Pt(2)

    p_ = doc.add_paragraph()
    run_num = p_.add_run(f'{num}  ')
    run_num.font.size = Pt(20)
    run_num.font.color.rgb = GOLD
    run_num.bold = True
    run_num.font.name = 'Microsoft YaHei'
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    run_title = p_.add_run(title)
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = INK
    run_title.bold = True
    run_title.font.name = 'Microsoft YaHei'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_after = Pt(4)

    # Gold rule (paragraph border bottom)
    pPr = p_._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'B8860B')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_subhead(text):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = SLATE
    run.italic = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_after = Pt(8)

def add_h2(text):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = INK
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_before = Pt(8)
    p_.paragraph_format.space_after = Pt(3)

def add_para(text, color=None, size=None, bold=False, italic=False):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if bold: run.bold = True
    if italic: run.italic = True
    return p_

def add_bullet(text):
    p_ = doc.add_paragraph()
    p_.paragraph_format.left_indent = Cm(0.5)
    run = p_.add_run(f'·  {text}')
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_after = Pt(2)
    return p_

def add_quote(text):
    p_ = doc.add_paragraph()
    p_.paragraph_format.left_indent = Cm(0.8)
    p_.paragraph_format.right_indent = Cm(0.8)
    # left gold bar via border
    pPr = p_._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'B8860B')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p_.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = INK
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_before = Pt(6)
    p_.paragraph_format.space_after = Pt(8)
    return p_

def add_image(path, width_cm):
    p_ = doc.add_paragraph()
    run = p_.add_run()
    run.add_picture(path, width=Cm(width_cm))
    p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p_

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, header_dark=True):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    # Set column widths later if needed
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p_ = cell.paragraphs[0]
        run = p_.add_run(h)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        if header_dark:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(cell, BG_DARK)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i+1].cells[c_i]
            cell.text = ''
            p_ = cell.paragraphs[0]
            run = p_.add_run(val)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(9.5)
            shade_cell(cell, BG_LIGHT)
    return t

def add_page_break():
    doc.add_page_break()

# ======================== TITLE PAGE ========================
# Big editorial title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('雅思自学')
run.font.size = Pt(60); run.bold = True
run.font.color.rgb = INK
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

title2 = doc.add_paragraph()
run2 = title2.add_run('从裸考一次开始')
run2.font.size = Pt(36); run2.italic = True
run2.font.color.rgb = INK
run2.font.name = 'Microsoft YaHei'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
eyebrow = doc.add_paragraph()
run = eyebrow.add_run('IELTS · SELF-STUDY PLAYBOOK')
run.font.size = Pt(10); run.bold = True
run.font.color.rgb = GOLD
run.font.name = 'Calibri'

sub = doc.add_paragraph()
run = sub.add_run('一份 66 分钟视频的全套方法论')
run.font.size = Pt(16)
run.font.color.rgb = SLATE
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
doc.add_paragraph()

# Three metadata cards
meta_card = doc.add_table(rows=1, cols=3)
meta_card.autofit = True
labels = [('TIME', '13 小时 / 日'), ('METHOD', '精听 · 语法 · 重复'), ('SOURCE', 'B 站 · Kasa_ZYY')]
for i, (k, v) in enumerate(labels):
    cell = meta_card.rows[0].cells[i]
    cell.text = ''
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(k)
    r1.font.size = Pt(8); r1.bold = True
    r1.font.color.rgb = GOLD
    r1.font.name = 'Calibri'
    p2 = cell.add_paragraph()
    r2 = p2.add_run(v)
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK
    r2.font.name = 'Microsoft YaHei'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
run = footer.add_run('整理：MiniMax-M3  ·  转录：Whisper (GPU) + 人工精读')
run.font.size = Pt(9)
run.font.color.rgb = SLATE
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

add_page_break()

# ======================== 01 · VIDEO INFO ========================
add_h1('01', '视频基本信息', 'At a Glance')

t = doc.add_table(rows=len([
    ['标题', '雅思自学流程介绍，以及一些奇技淫巧'],
    ['UP 主', 'Kasa_ZYY'],
    ['BV 号', 'BV1cyDKBLEXY'],
    ['时长', '66 分 15 秒（3975 秒）'],
    ['发布日期', '2026-04-04'],
    ['数据', '318,541 播放 · 22,719 赞 · 48,955 收藏 · 726 弹幕'],
    ['分区', '学习 · 英语 · 雅思'],
    ['标签', '自学 / 英语 / 雅思 / 学习 / 教程 / IELTS'],
]), cols=2)
t.style = 'Table Grid'
for r_i, (k, v) in enumerate([
    ['标题', '雅思自学流程介绍，以及一些奇技淫巧'],
    ['UP 主', 'Kasa_ZYY'],
    ['BV 号', 'BV1cyDKBLEXY'],
    ['时长', '66 分 15 秒（3975 秒）'],
    ['发布日期', '2026-04-04'],
    ['数据', '318,541 播放 · 22,719 赞 · 48,955 收藏 · 726 弹幕'],
    ['分区', '学习 · 英语 · 雅思'],
    ['标签', '自学 / 英语 / 雅思 / 学习 / 教程 / IELTS'],
]):
    cell_k = t.rows[r_i].cells[0]
    cell_v = t.rows[r_i].cells[1]
    cell_k.text = ''
    p1 = cell_k.paragraphs[0]
    r1 = p1.add_run(k)
    r1.font.size = Pt(9); r1.bold = True
    r1.font.color.rgb = GOLD
    r1.font.name = 'Microsoft YaHei'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(cell_k, BG_BAND)
    cell_v.text = ''
    p2 = cell_v.paragraphs[0]
    r2 = p2.add_run(v)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    r2.font.name = 'Microsoft YaHei'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(cell_v, BG_LIGHT)

doc.add_paragraph()
add_para('内容主线：一份完整的雅思自学日常时间表，加上听力/阅读/写作/词汇的具体方法，以及作者亲身验证过的若干「奇技淫巧」。')

add_page_break()

# ======================== 02 · TL;DR INFOGRAPHIC ========================
add_h1('02', '七条核心结论', 'TL;DR')
add_subhead('60 秒读完 · 详见后续章节展开')

add_image(r'C:\Users\Administrator\Desktop\雅思自学-TL;DR-信息图.png', 17)

add_page_break()

# ======================== 03 · SCHEDULE ========================
add_h1('03', '每日作息', 'Daily Schedule')
add_subhead('作者本人 13 小时备考节奏 · 三个时段、三种能量')

add_image(r'C:\Users\Administrator\Desktop\diagrams\daily-schedule.png', 17)
doc.add_paragraph()

# Quadrant commentary — proper 2x2 grid, 1 quadrant per cell (header + content)
quad_t = doc.add_table(rows=2, cols=2)
quad_t.style = 'Table Grid'

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)

quadrants = [
    # (col, row, color, header, bullets)
    (0, 0, PURPLE,  '上午 · 输入',
     ['词汇是 4 项技能的底座',
      '100-150 新词 + 复习 150-200']),
    (1, 0, EMERALD, '下午 · 练习',
     ['完整一套机考（听力 + 阅读）',
      '报考时段 = 练习时段']),
    (0, 1, AMBER,   '晚上 · 输出',
     ['写作早开始，写完 AI 改',
      '模板不动，只升一档']),
    (1, 1, SLATE,   '休息 · 切换',
     ['午睡 5 分钟提神',
      '晚饭彻底离开学习环境']),
]

for col, row, color, header, bullets in quadrants:
    cell = quad_t.rows[row].cells[col]
    cell.text = ''
    # Header
    p_ = cell.paragraphs[0]
    run = p_.add_run(header)
    run.font.size = Pt(12); run.bold = True
    run.font.color.rgb = color
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_.paragraph_format.space_after = Pt(4)
    # Bullets
    for b in bullets:
        p_ = cell.add_paragraph()
        p_.paragraph_format.left_indent = Cm(0.4)
        run = p_.add_run(f'·  {b}')
        run.font.size = Pt(9.5)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p_.paragraph_format.space_after = Pt(3)
    shade_cell(cell, BG_LIGHT)

add_page_break()

# ======================== 04 · LISTENING ========================
add_h1('04', '听力 · 奇技淫巧集中地', 'Listening')
add_subhead('考场节奏 · 精听流程 · 错题归档 · 基础信息')

add_image(r'C:\Users\Administrator\Desktop\diagrams\intensive-listening-flow.png', 15)
doc.add_paragraph()

add_h2('4.1  考场节奏 · 平时练习的核心技巧')
for t_ in [
    '选项要「一眼翻译」：训练自己一秒钟内把匹配题选项的中文意思直接翻译出来。',
    '录音开始放寒暄时（"This is the British Council..."），立刻翻到 P3 先读题。',
    'P2/P3 比 P1/P4 难，重点突破匹配题。',
    'P1 是填词题多，但要警惕偶发的选择题（剑 7/8 出现过，作者亲历）。',
    '做完 P1 之后，你比其他人多看了 P2 一遍题。',
    '填词题：单词、冠词、单复数、时态、大小写一个都不能错。',
]:
    add_bullet(t_)

add_h2('4.2  复盘 · 精听（最最关键）')
add_quote('精听是提高听力的唯一方法。不要全文精听，要单句精听。')
for step in [
    '听 1 遍 → 2 遍 → 3 遍',
    '仍不懂 → 打开译文，对照听',
    '还不懂 → 把句子丢给 AI 做语法分析：主谓宾 / 状语 / 定语 / 从句',
    '真正读懂结构 → 下次听到同结构立即反应',
]:
    add_bullet(step)

add_h2('4.3  错题归档（用飞书 / Excel）')
for t_ in [
    '精听完开始改错，把错误分类：拼写 / 大小写 / 单复数 / 单词不会',
    '所有错词加入每日早上的单词背诵流程',
    '不要追求「精听会不会太简单所以不听」——只要一遍就懂，就过',
]:
    add_bullet(t_)

add_h2('4.4  地图题')
for t_ in [
    '本质只是方位词：东南西北 / 红绿灯 / 十字路口',
    '把这些方位词读懂，地图题就很简单',
    '听的时候跟着图形一格一格走，不要提前猜',
]:
    add_bullet(t_)

add_h2('4.5  必须敏感的基础信息')
for t_ in [
    '数字：电话号码、护照号、信用卡号等长串要瞬间反应',
    '月份拼写：January 不是 Jan first，是 January first',
    '星期、颜色、国家（Portuguese 等）必须反映迅速',
    'm vs n：人名地名发音（Wombat ≠ Wonbat）第一遍就要听完整名字',
]:
    add_bullet(t_)

add_page_break()

# ======================== 05 · VOCAB ========================
add_h1('05', '词汇 · 一切的地基', 'Vocabulary')
add_subhead('原则、拼写、流程、不要做的事')

add_h2('5.1  核心原则')
for t_ in [
    '背单词本质是「重复，重复，还是重复」。',
    '每天 100-150 新词，复习 150-200，绝对不能含糊。',
    '三秒规则：3 秒内想不出中文意思 = 不认识，立刻标「不认识」。不要骗自己。',
    '「四会非会就是不会，四懂非懂就是不懂，不懂装懂那更是不懂。」',
    '发音 + 拼写都要记：听力考发音，写作/填空题考拼写。',
]:
    add_bullet(t_)

add_h2('5.2  拼写就是分数')
add_quote('一个拼写错误 = 1 分没拿到。雅思听力 40 题 / 40 分，错 1 个就是 29 → 6.5 分，跟 7 分天差地别。')

add_h2('5.3  流程演示（基于 App）')
for t_ in [
    '背 100 个新词 → 完成',
    '次日：再学新 100 词 + 复习昨天',
    '背到 500 词以后：每天整体过 500 词',
    '背到 1000 词以后：每天整体过 1000 词',
    '三秒想不出 → 立刻标错 → 加入「不认识」清单 → 之后再重点背',
]:
    add_bullet(t_)

add_h2('5.4  不要做的事')
for t_ in [
    '不要花大量时间抠「单词意识」/ 字眼，只记最常用的中文对应即可。',
    '真的完全查不到的中文意思（如 turbulence 生僻义），再去查 AI，不要逐字查。',
    '不要在阅读时把每个生词都记下来——只记影响你理解句子主干的单词。',
]:
    add_bullet(t_)

add_page_break()

# ======================== 06 · READING ========================
add_h1('06', '阅读', 'Reading')
add_subhead('题型策略 · 复盘动作')

add_h2('6.1  题型策略')
t = doc.add_table(rows=5, cols=2)
t.style = 'Table Grid'
header = t.rows[0].cells
for i, h in enumerate(['题型', '策略']):
    header[i].text = ''
    p_ = header[i].paragraphs[0]
    run = p_.add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(header[i], BG_DARK)
for i, (k, v) in enumerate([
    ('填空题', '优先圈出关键定位词（人名 / 时间 / 数字）'),
    ('判断题 T/F/NG', '读 2 道题 → 读 1 段 → 对应判断；拿不准果断跳过'),
    ('匹配题', '不要靠一句话定位，理解整段逻辑主干'),
    ('单选题', '问题一定要读懂，选项逐个搞懂；错了就做语法分析'),
]):
    cells = t.rows[i+1].cells
    for j, val in enumerate([k, v]):
        cells[j].text = ''
        p_ = cells[j].paragraphs[0]
        run = p_.add_run(val)
        run.font.size = Pt(9.5)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        shade_cell(cells[j], BG_LIGHT)

add_h2('6.2  复盘 · 关键动作')
for t_ in [
    '不要全文标注生词（感动自己），只标影响理解主干的词',
    '不懂的句子 → 复制粘贴给 AI 做语法分析，搞清结构',
    '语法分析不是每句都做：明显简单的不用；不简单的才做',
    '不要迷信「定位」：IELTS 越来越难，本质上还是考对文章逻辑的理解',
    '做完所有题后回头看：定位 + 句子理解双线并行',
]:
    add_bullet(t_)

add_page_break()

# ======================== 07 · WRITING ========================
add_h1('07', '写作', 'Writing')
add_subhead('越早开始 · AI 模板升级 · 错误自查')

add_h2('7.1  总体策略')
for t_ in [
    '越早开始写越好：水平差也要硬写，错误是改出来的',
    '大作文建议写 400 词左右（要求 250，多写 ≈ 高分）',
    '小作文 150-200 词，按题型准备模板',
    '写作 6.5 = 语法好 + 模板稳，并不需要多华丽的单词',
]:
    add_bullet(t_)

add_h2('7.2  AI 改作文的正确打开方式')
add_quote('让 AI 在你的模板基础上改，只升级到下一档分数（如 6 → 7），不要直接让 AI 给你写一篇 9 分。')
for s in [
    '写完一篇',
    '丢给 AI：「请基于我的水平、保持我的模板结构，把用词 / 句式升级到 X.X 分」',
    '拿 AI 改后的版本与自己的版本做对比',
    '逐句问 AI：「为什么改成这样就能到 7 分？」',
    '如此循环，写作能力会真正提升',
]:
    add_bullet(s)

add_h2('7.3  必须自查的错误类型')
for t_ in ['主谓一致 (subject-verb agreement)', '单复数', '拼写 / 大小写', '时态', '指代不明']:
    add_bullet(t_)

add_page_break()

# ======================== 08 · SPEAKING ========================
add_h1('08', '口语', 'Speaking')
add_subhead('写作好 → 口语不会差 · 研究题库 ≠ 背答案')

for t_ in [
    '作者本人口语一直 6 分，但他认为「写作好 → 口语不会差」，二者底层能力相通',
    '如果学校卡小分（6.5 / 7），需要专门研究口语题库',
    '研究题库 ≠ 背答案：研究的是「评分标准在意的格式、剧情、表达方式」',
    '找专业老师做 1-2 次模拟：让老师评价当前分段 + 给提升路径',
    '不要每周都练，1-2 次专项即可',
]:
    add_bullet(t_)

add_page_break()

# ======================== 09 · PRO TIPS ========================
add_h1('09', '临场技巧 · 奇技淫巧', 'Pro Tips')

add_h2('9.1  听力耳机音量（独家技巧）')
for t_ in [
    '练习时音量保持 70% 左右',
    '考试那天调到 100% — 平时 70 已经 OK 的，到考场 100 会让你觉得「哇，好清晰」',
    '如果你平时就用 100%，考试就没法再升一级了',
]:
    add_bullet(t_)

add_h2('9.2  考试开始前的「环境观察」')
for t_ in [
    '进入机考页面后会先播放考试须知',
    '放完后会给你一个「播放按钮」——不要立即点！',
    '先观察周围同学：是否有老师正在帮同学处理登录问题？',
    '等全场同学都开始放听力了，老师也走了，再点开始',
    '只要不点，时间就不算你的听力时间',
]:
    add_bullet(t_)

add_h2('9.3  考试中的小动作')
for t_ in [
    '可以带水进考场（作者经验允许），紧张时含一口',
    '不要喝多，喝多要去厕所 → 浪费时间，且时间不暂停',
]:
    add_bullet(t_)

add_h2('9.4  报考时段选择')
add_quote('练习在哪个时段，就报哪个时段的考试。大脑会在固定时段自动进入「考试模式」。')

add_h2('9.5  剑雅真题使用顺序')
for t_ in [
    '建议先做 剑 20 → 剑 18：这些版本前面会放 "this is the British Council..." 的寒暄，给你足够时间练「先读 P3」',
    '剑 17/18 之后开始没有寒暄了，直接进 P1，时间紧',
    '所以先用老版本练节奏，再做新版本',
    '老版本（剑 7/8/9）也要做：题型有时会复古',
    '做过老题 → 考场上不慌',
]:
    add_bullet(t_)

add_page_break()

# ======================== 10 · MONEY MAP ========================
add_h1('10', '钱花在哪里 · 不花在哪里', 'Money Map')

t = doc.add_table(rows=9, cols=3)
t.style = 'Table Grid'
header = t.rows[0].cells
for i, h in enumerate(['建议', '项目', '理由']):
    header[i].text = ''
    p_ = header[i].paragraphs[0]
    run = p_.add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(header[i], BG_DARK)

money_rows = [
    ('推荐', EMERALD, 'AI 会员（ChatGPT 等）', '反馈、情绪价值、持续激励、未来通用技能'),
    ('推荐', EMERALD, '专业老师 1 对 1 专项方法课', '当你卡在某一分段很久，需要人打破思维僵局'),
    ('推荐', EMERALD, '写作老师人工批改', '与 AI 互补，更精准、更便宜'),
    ('推荐', EMERALD, '口语模拟 1-2 次', '让老师给出明确的提升路径与目标分要求'),
    ('不建议', RED, '雅思全程班 / 冲刺班', '信息透明、AI 已能覆盖；动辄上万，本质靠自己'),
    ('不建议', RED, '「3 天 7 天速成」课程', '焦虑营销、装逼话术、毫无玄学'),
    ('不建议', RED, '真题库（不是剑雅）会员', '碎片化、缺连贯节奏；不如把剑雅做三遍'),
    ('看情况', AMBER, '心理咨询', '当你焦虑、心态崩时比报班有用得多'),
]
for i, (status, color, item, reason) in enumerate(money_rows):
    cells = t.rows[i+1].cells
    cells[0].text = ''
    p_ = cells[0].paragraphs[0]
    run = p_.add_run(status)
    run.font.size = Pt(9); run.bold = True
    run.font.color.rgb = color
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    cells[1].text = ''
    p_ = cells[1].paragraphs[0]
    run = p_.add_run(item)
    run.font.size = Pt(9.5)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    cells[2].text = ''
    p_ = cells[2].paragraphs[0]
    run = p_.add_run(reason)
    run.font.size = Pt(9.5)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(cells[0], BG_BAND)
    shade_cell(cells[1], BG_LIGHT)
    shade_cell(cells[2], BG_LIGHT)

add_page_break()

# ======================== 11 · KEY QUOTES ========================
add_h1('11', '关键金句', 'Key Quotes')
add_subhead('— 直接摘自原视频，建议贴在书桌前')

for q in [
    '如果你语法都不懂，赶紧停下来，先去学语法。',
    '四会非会就是不会，四懂非懂就是不懂，不懂装懂那更是不懂。',
    '环境是真的非常重要。环境会倒逼自己自学。',
    '雅思这么多年了，技巧性的东西在网上都能找到属于自己的方法。',
    '花 2000 块去考一次，不是为了过，是为了熟悉流程 + 让自己心痛。',
    'AI 不仅是工具，还是情绪价值提供者。',
    '语法真的是一切的基础。',
    '报名时段跟着练习时段走，让大脑自动进入考试状态。',
    '报考下午场就下午练，报考上午就上午练。',
    '听力复习唯一的方法：精听，单句精听。',
    'AI 改作文要在你模板上改到下一档，不要直接写 9 分给你。',
    '小红书 / 抖音删掉，那上面的「8 分有手就行」会影响你心态。',
]:
    add_quote(q)

add_page_break()

# ======================== 12 · TIMESTAMP INDEX ========================
add_h1('12', '附录 · 时间戳索引', 'Appendix')
add_para('以下时间戳来自原视频，可对照 B 站跳转到对应段落。', color=SLATE, size=9)
doc.add_paragraph()

ts = [
    ('00:00 - 05:00', '开场建议：机考训练 / 第一次裸考 / 报班 vs 自学 / 环境'),
    ('05:00 - 10:00', '钱花在哪：AI 会员 / 何时找老师 / 写作口语一对一'),
    ('10:00 - 15:00', '真题/速成课避坑 + 每日时间表 + 不碰手机原则'),
    ('15:00 - 20:00', '背单词方法：100-150 新词 / 三秒规则 / 整体过'),
    ('20:00 - 25:00', '拼写重要性 / 听力选项翻译 / 精听入门'),
    ('25:00 - 30:00', '听力考场节奏 / 阅读题型策略 / 平行阅读'),
    ('30:00 - 35:00', '判断题节奏 / 精听 3 遍流程 / 语法是基础'),
    ('35:00 - 40:00', '错题归档 / 地图题 / 数字月份训练'),
    ('40:00 - 45:00', '国家名 / m-n 区分 / 阅读复盘 + AI 语法分析'),
    ('45:00 - 50:00', '不做全文标注 / 匹配题逻辑主干 / 语法分析时机'),
    ('50:00 - 55:00', '判断题问 AI / 写作要早开始 / 模板升级策略'),
    ('55:00 - 60:00', '口语题库研究 / 耳机音量 70 技巧 / 考试环境观察'),
    ('60:00 - 65:00', '删小红书抖音 / 早睡早起 / 报下午场'),
    ('65:00 - 66:15', '剑雅 7/8/9 老题价值 / 收尾'),
]
t = doc.add_table(rows=1+len(ts), cols=2)
t.style = 'Table Grid'
header = t.rows[0].cells
for i, h in enumerate(['时段', '内容']):
    header[i].text = ''
    p_ = header[i].paragraphs[0]
    run = p_.add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(header[i], BG_DARK)
for i, (k, v) in enumerate(ts):
    cells = t.rows[i+1].cells
    cells[0].text = ''
    p_ = cells[0].paragraphs[0]
    run = p_.add_run(k)
    run.font.size = Pt(9.5)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(cells[0], BG_BAND)
    cells[1].text = ''
    p_ = cells[1].paragraphs[0]
    run = p_.add_run(v)
    run.font.size = Pt(9.5)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade_cell(cells[1], BG_LIGHT)

doc.add_paragraph()
doc.add_paragraph()
add_para('※ 本总结基于 Whisper (small 模型, GPU) 转录后人工精读整理。', color=SLATE, size=8)
add_para('※ 如需引用本文内容，请注明原视频出处：B 站 BV1cyDKBLEXY · UP 主 Kasa_ZYY', color=SLATE, size=8)

doc.save(OUTPUT)
import os
print('Docx written to:', OUTPUT)
print('Size:', os.path.getsize(OUTPUT), 'bytes')