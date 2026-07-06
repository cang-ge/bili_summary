"""
DOCX renderer — driven by `data['layout']` built by llm_client._build_layout().

Each layout item is dispatched by its `type` field to the matching section
builder.  Sections with no meaningful content are absent from the layout
array and thus never rendered.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from design_tokens import (
    INK_DOCX as INK,
    GOLD_DOCX as GOLD,
    PAPER_DOCX as PAPER,
    SLATE_DOCX as SLATE,
    BORDER_DOCX as BORDER,
    EMERALD_DOCX as EMERALD,
    AMBER_DOCX as AMBER,
    RED_DOCX as RED,
    BG_DARK,
    BG_LIGHT,
    BG_BAND,
)

FONT = 'Microsoft YaHei'


def _set_run_font(run, size=None, color=None, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _h1(doc, num, title, eyebrow=None):
    if eyebrow:
        p_ = doc.add_paragraph()
        run = p_.add_run(eyebrow.upper())
        _set_run_font(run, size=9, color=GOLD, bold=True)
        p_.paragraph_format.space_after = Pt(2)
    p_ = doc.add_paragraph()
    if num:
        run_num = p_.add_run(f'{num}  ')
        _set_run_font(run_num, size=20, color=GOLD, bold=True)
    run_title = p_.add_run(title)
    _set_run_font(run_title, size=22, color=INK, bold=True)
    p_.paragraph_format.space_after = Pt(4)
    # Gold rule (paragraph border bottom)
    pPr = p_._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'B8860B')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, text):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    _set_run_font(run, size=13, color=INK, bold=True)
    p_.paragraph_format.space_before = Pt(8)
    p_.paragraph_format.space_after = Pt(3)


def _h3(doc, text):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    _set_run_font(run, size=11, color=SLATE, italic=True)
    p_.paragraph_format.space_before = Pt(6)
    p_.paragraph_format.space_after = Pt(2)


def _subhead(doc, text):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    _set_run_font(run, size=11, color=SLATE, italic=True)
    p_.paragraph_format.space_after = Pt(8)


def _para(doc, text, *, size=10, color=INK, italic=False, bold=False):
    p_ = doc.add_paragraph()
    run = p_.add_run(text)
    _set_run_font(run, size=size, color=color, italic=italic, bold=bold)
    p_.paragraph_format.space_after = Pt(4)
    return p_


def _bullet(doc, text):
    p_ = doc.add_paragraph()
    p_.paragraph_format.left_indent = Cm(0.5)
    run = p_.add_run(f'·  {text}')
    _set_run_font(run, size=10)
    p_.paragraph_format.space_after = Pt(2)
    return p_


def _quote(doc, text):
    p_ = doc.add_paragraph()
    p_.paragraph_format.left_indent = Cm(0.8)
    p_.paragraph_format.right_indent = Cm(0.8)
    p_.paragraph_format.space_before = Pt(6)
    p_.paragraph_format.space_after = Pt(6)
    run_bar = p_.add_run('▍ ')
    _set_run_font(run_bar, size=11, color=GOLD, bold=True)
    run_text = p_.add_run(text)
    _set_run_font(run_text, size=10.5, color=INK, italic=True)
    # Left gold bar via paragraph border
    pPr = p_._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'B8860B')
    pBdr.append(left)
    pPr.append(pBdr)


def _title_page(doc, cover_data, meta_data):
    meta = meta_data or {}
    cover = cover_data or {}
    title = cover.get('display_title') or meta.get('title') or '视频总结'
    subtitle = cover.get('subtitle') or ''
    tagline = cover.get('tagline') or ''
    duration = meta.get('duration_label') or ''

    # Big title
    p_ = doc.add_paragraph()
    p_.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_.add_run(title[:30])
    _set_run_font(run, size=60, color=INK, bold=True)
    p_.paragraph_format.space_after = Pt(8)

    if subtitle:
        p_ = doc.add_paragraph()
        run = p_.add_run(subtitle[:40])
        _set_run_font(run, size=24, color=INK, italic=True)
        p_.paragraph_format.space_after = Pt(6)

    p_ = doc.add_paragraph()
    run = p_.add_run('B 站视频 · 自动生成总结')
    _set_run_font(run, size=10, color=GOLD, bold=True)
    p_.paragraph_format.space_after = Pt(8)

    if tagline:
        p_ = doc.add_paragraph()
        run = p_.add_run(tagline[:80])
        _set_run_font(run, size=12, color=SLATE)
        p_.paragraph_format.space_after = Pt(16)

    # Meta card row
    cards = [(meta.get('uploader') or '—', 'UP 主'),
             (duration or '—', '时长'),
             (meta.get('view') or '—', '播放量')]
    table = doc.add_table(rows=1, cols=3)
    for cell, (val, label) in zip(table.rows[0].cells, cards):
        cell.text = ''
        p1 = cell.paragraphs[0]
        run = p1.add_run(val)
        _set_run_font(run, size=14, color=INK, bold=True)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = cell.add_paragraph()
        run = p2.add_run(label)
        _set_run_font(run, size=9, color=GOLD, bold=True)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def _meta_table(doc, meta, eyebrow, en_eyebrow):
    _h1(doc, eyebrow, '视频基本信息', en_eyebrow)

    rows = [
        ('标题', meta.get('title') or '—'),
        ('UP 主', meta.get('uploader') or '—'),
        ('BV 号', meta.get('bvid') or '—'),
        ('时长', meta.get('duration_label') or '—'),
        ('发布日期', str(meta.get('pubdate') or '—')),
    ]
    data_line = ' · '.join(filter(None, [
        f"{meta.get('view')} 播放" if meta.get('view') else '',
        f"{meta.get('like')} 赞" if meta.get('like') else '',
        f"{meta.get('favorite')} 收藏" if meta.get('favorite') else '',
        f"{meta.get('danmaku')} 弹幕" if meta.get('danmaku') else '',
    ]))
    if data_line:
        rows.append(('数据', data_line))
    if meta.get('tname'):
        rows.append(('分区', meta['tname']))
    if meta.get('tags'):
        rows.append(('标签', ' / '.join(meta['tags'])))
    rows = [r for r in rows if r[1] and r[1] != '—']

    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for row, (k, v) in zip(table.rows, rows):
        row.cells[0].text = ''
        p_ = row.cells[0].paragraphs[0]
        run = p_.add_run(k)
        _set_run_font(run, size=10, color=GOLD, bold=True)
        row.cells[1].text = ''
        p_ = row.cells[1].paragraphs[0]
        run = p_.add_run(v)
        _set_run_font(run, size=10, color=INK)

    doc.add_page_break()


def _infographic_page(doc, assets, info):
    info = info if isinstance(info, dict) else {}
    eyebrow = (info.get('eyebrow') or 'CORE TAKEAWAYS').upper()
    hero = info.get('hero_title') or '核心结论'

    p_ = doc.add_paragraph()
    run = p_.add_run(eyebrow)
    _set_run_font(run, size=9, color=GOLD, bold=True)
    p_.paragraph_format.space_after = Pt(2)

    p_ = doc.add_paragraph()
    run = p_.add_run(hero)
    _set_run_font(run, size=28, color=INK, bold=True)
    p_.paragraph_format.space_after = Pt(2)

    p_ = doc.add_paragraph()
    run = p_.add_run('TL;DR · 一图速览')
    _set_run_font(run, size=12, color=SLATE, italic=True)
    p_.paragraph_format.space_after = Pt(8)

    info_path = assets.get('infographic')
    if info_path and Path(info_path).exists():
        p_ = doc.add_paragraph()
        run = p_.add_run()
        run.add_picture(str(info_path), width=Cm(16))
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def _section_chapter(doc, d, eyebrow, en_eyebrow):
    title = d.get('title') or '章节'
    _h1(doc, eyebrow, title, en_eyebrow)
    if d.get('subtitle'):
        _subhead(doc, d['subtitle'])
    if d.get('body'):
        _para(doc, d['body'])
    for b in (d.get('bullets') or []):
        _bullet(doc, str(b))
    for sub in (d.get('subsections') or []):
        if sub.get('heading'):
            _h3(doc, sub['heading'])
        for b in (sub.get('bullets') or []):
            _bullet(doc, str(b))
    if d.get('pull_quote'):
        _quote(doc, d['pull_quote'])


def _pro_tips(doc, d, eyebrow, en_eyebrow):
    items = d.get('items') or []
    if not items:
        return
    doc.add_page_break()
    title = d.get('intro') or '临场技巧'
    _h1(doc, eyebrow, title, en_eyebrow)
    for it in items:
        if it.get('heading'):
            _h2(doc, it['heading'])
        for b in (it.get('bullets') or []):
            _bullet(doc, str(b))


def _money_map(doc, d, eyebrow, en_eyebrow):
    rows = d.get('rows') or []
    if len(rows) < 3:
        return
    doc.add_page_break()
    _h1(doc, eyebrow, d.get('intro') or '钱花在哪里 · 不花在哪里',
        en_eyebrow)

    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.autofit = False
    table.rows[0].cells[0].text = ''
    p_ = table.rows[0].cells[0].paragraphs[0]
    run = p_.add_run('建议')
    _set_run_font(run, size=10, color=GOLD, bold=True)
    for ci, h in enumerate(['项目', '理由'], start=1):
        table.rows[0].cells[ci].text = ''
        run = table.rows[0].cells[ci].paragraphs[0].add_run(h)
        _set_run_font(run, size=10, color=GOLD, bold=True)

    for i, r in enumerate(rows, start=1):
        for ci, key in enumerate(['status', 'item', 'reason']):
            table.rows[i].cells[ci].text = ''
            run = table.rows[i].cells[ci].paragraphs[0].add_run(r.get(key, ''))
            _set_run_font(run, size=10, color=INK)


def _quotes(doc, d, eyebrow, en_eyebrow):
    items = d.get('items') or []
    items = [q for q in items if q]
    if not items:
        return
    doc.add_page_break()
    _h1(doc, eyebrow, '关键金句', en_eyebrow)
    for q in items[:12]:
        _quote(doc, q)


def _timestamp_index(doc, d, eyebrow, en_eyebrow):
    rows = d.get('rows') or []
    if len(rows) < 3:
        return
    doc.add_page_break()
    _h1(doc, eyebrow, '附录 · 时间戳索引', en_eyebrow)
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.rows[0].cells[0].text = ''
    run = table.rows[0].cells[0].paragraphs[0].add_run('时间范围')
    _set_run_font(run, size=10, color=GOLD, bold=True)
    table.rows[0].cells[1].text = ''
    run = table.rows[0].cells[1].paragraphs[0].add_run('主题')
    _set_run_font(run, size=10, color=GOLD, bold=True)
    for i, r in enumerate(rows, start=1):
        for ci, key in enumerate(['range', 'topic']):
            table.rows[i].cells[ci].text = ''
            run = table.rows[i].cells[ci].paragraphs[0].add_run(r.get(key, ''))
            _set_run_font(run, size=10, color=INK)


def _closing(doc, closing):
    closing = closing if isinstance(closing, dict) else {}
    doc.add_paragraph()
    if closing.get('disclaimer'):
        p_ = doc.add_paragraph()
        run = p_.add_run(closing['disclaimer'])
        _set_run_font(run, size=8, color=SLATE)
    if closing.get('credit'):
        p_ = doc.add_paragraph()
        run = p_.add_run(closing['credit'])
        _set_run_font(run, size=8, color=SLATE)


def render_docx(data: dict, out_path: Path) -> None:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # Default font
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    assets = data.get('_assets') or {}
    meta = data.get('meta') or {}

    # ---- Layout-driven rendering loop ----
    cover_data = None
    layout = data.get('layout') or []
    for item in layout:
        t = item['type']
        d = item['data']
        eb = item.get('eyebrow', '')
        en_eb = item.get('en_eyebrow', '')

        if t == 'cover':
            cover_data = d
        elif t == 'meta':
            _title_page(doc, cover_data, meta)
            _meta_table(doc, d, eb, en_eb)
        elif t == 'chapter':
            _section_chapter(doc, d, eb, en_eb)
        elif t == 'pro_tips':
            _pro_tips(doc, d, eb, en_eb)
        elif t == 'money_map':
            _money_map(doc, d, eb, en_eb)
        elif t == 'quotes':
            _quotes(doc, d, eb, en_eb)
        elif t == 'infographic':
            if info_path := assets.get('infographic'):
                _infographic_page(doc, assets, d)
        elif t == 'timestamp_index':
            _timestamp_index(doc, d, eb, en_eb)
        elif t == 'closing':
            _closing(doc, d)

    doc.save(str(out_path))
    print(f'DOCX saved: {out_path} ({out_path.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    sample = {
        'meta': {
            'title': '测试视频', 'uploader': 'TestUser', 'bvid': 'BV1xxxx',
            'pubdate': '2026-07-06', 'duration_label': '10 分',
            'view': '1.0 万', 'like': '500', 'favorite': '200', 'danmaku': '30',
            'tname': '知识', 'tags': ['测试', 'demo'],
        },
        'cover': {'display_title': '测试视频', 'subtitle': '副标题',
                  'tagline': '一句话介绍'},
        'intro': {'paragraph': '这是一个测试视频的引言。'},
        'layout': [
            {'type': 'cover', 'data': {'display_title': '测试视频', 'subtitle': '副标题', 'tagline': '一句话介绍'}},
            {'type': 'meta', 'data': {
                'title': '测试视频', 'uploader': 'TestUser', 'bvid': 'BV1xxxx',
                'pubdate': '2026-07-06', 'duration_label': '10 分',
                'view': '1.0 万', 'like': '500', 'favorite': '200', 'danmaku': '30',
                'tname': '知识', 'tags': ['测试', 'demo'],
            }, 'eyebrow': '01', 'en_eyebrow': 'At a Glance'},
            {'type': 'intro', 'data': {'paragraph': '这是一个测试视频的引言。'}},
            {'type': 'chapter', 'data': {'title': '第一章节', 'body': '章节正文。',
             'bullets': ['要点 1', '要点 2'], 'pull_quote': '金句示例。'},
             'eyebrow': '02', 'en_eyebrow': 'Chapter 1'},
            {'type': 'chapter', 'data': {'title': '第二章节', 'body': '继续。',
             'bullets': ['a', 'b']}, 'eyebrow': '03'},
            {'type': 'quotes', 'data': {'items': ['第一句', '第二句']},
             'eyebrow': 'QUOTES', 'en_eyebrow': 'Key Quotes'},
            {'type': 'closing', 'data': {'disclaimer': '※ 自动生成的总结', 'credit': '整理 · TestUser'}},
        ],
    }
    render_docx(sample, Path('output/summary.docx'))